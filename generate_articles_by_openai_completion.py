# coding: utf-8

import csv
import time
import requests
import os
import openai
from docx import Document
from loguru import logger
from pkg import parse_curl, utils, lark_bot
from typing import List


class Client(object):
    def __init__(self) -> None:
        pass


def raise_on_unexpected_resp(resp):
    if resp.status_code != 200:
        status_code, content = resp.status_code, resp.content
        err_msg = f'unexpected response, status_code={status_code}, content={content}'
        logger.error(err_msg)
        raise Exception(err_msg)


def generate_article_by_title(title: str) -> dict:
    logger.info(f'ask openai write content about `{title}`')

    response = openai.Completion.create(
        model="text-davinci-003",
        prompt=f"{title}，写一篇350字文章",
        max_tokens=1500,
    )
    article = response.choices[0].text
    return article
    

def load_tasks() -> List[dict]:
    with open('output/tasks.csv', 'r', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        return [item for item in reader]


def generate_article_docx(article_data: dict, output_path: str):
    document = Document()
    document.add_heading(article_data['title'], 0)

    # 增加介绍
    # document.add_paragraph(article_data['intro'])

    # 写入正文
    document.add_paragraph(article_data['content'])
    # contents = article_data['contents']
    # for content in contents:
    #     document.add_heading(content['title'], level=2)
    #     document.add_paragraph(content['intro'])

    document.save(output_path)

def main2():
    article = generate_article_by_title('关于二婚的看法')
    print(article)


def is_contains_words(s: str, words: List[str]) -> bool:
    for word in words:
        if s.find(word) != -1:
            return True
    return False


@lark_bot.exception_notify
def main():
    # 加载配置文件
    settings = utils.load_json_from_file('conf/settings.json')
    stop = 0
    # 配置 lark bot 
    lark_bot_webhook = settings.get('lark_bot_webhook')
    if lark_bot_webhook:
        lark_bot.init_with_webhook(lark_bot_webhook)

    # 配置 openai api key
    openapi_api_key = settings.get('openai', {}).get('api_key')
    if not openapi_api_key:
        raise Exception('为设置 openai api key')
    openai.api_key = openapi_api_key

    # 加载任务列表
    tasks = load_tasks()
    logger.info(f'{len(tasks)} tasks is loaded')

    
    ignore_words = settings['ignore_words']

    # 遍历任务列表生成文章
    for task in tasks:
        title = task['title']
        category = task['category']

        if len(title) < 5:
            # 标题长度不得小于5个字
            continue

        if is_contains_words(title, ignore_words):
            # 忽略标题含有某些字或字段
            continue
    
        output_dir = f'output/articles/{category}'
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        output_path = f'{output_dir}/{title}.docx'
        if os.path.exists(output_path):
            logger.info(f'`{output_path}` already exists')
            continue
        
        logger.info(f'generate article by title, title={title}')
        error = 0
        start_at = time.time()
        while True:
            try:
                article_content = generate_article_by_title(title)
                article_data = {
                    'title': title,
                    'content': article_content,
                }
                elapsed = time.time() - start_at
                logger.info(f'generate article by title success, title={title}, cost={elapsed:.2f}s')
                logger.info(f'generate article docx by title, title={title}')
                generate_article_docx(article_data, output_path)
            except OSError:
                break
            except Exception as e:
                if str(e) == 'The server is overloaded or not ready yet.':
                    logger.info(f'服务器繁忙，正在重试')
                    continue 
                if str(e) == 'You exceeded your current quota, please check your plan and billing details.':
                    logger.inf(f'请更换api_key')
                    stop = 1
                    break
                else:
                    error += 1
                    print('出现未知错误,正在重试')
                    print(e)
                    if error > 20:
                        stop = 1
                        break
            else:
                break

        if stop == 1:
            print('api_key用尽或出现其它错误,请检查')
            break


if __name__ == '__main__':
    main()
